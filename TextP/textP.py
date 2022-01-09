import tkinter as tk
import tkinter.font as tkFont
from tkinter.constants import BOTH, TOP, YES

from PIL import Image, ImageTk
import pyttsx3  

from tkinter.filedialog import askopenfile
import docx

import pathlib
import clipboard
import speech_recognition as sr 
  
# Initialize the recognizer  
r = sr.Recognizer()  

root = tk.Tk()
def show_frame(frame):
    frame.tkraise()
    

root.title("TextP")
root.iconbitmap(r,'TextP_icon.ico')
root.geometry("900x600")

root.resizable(False, False)

root.configure(bg='#F0F0F0')

#canvas = tk.Canvas(root, width=800, height=500)
#canvas.grid(columnspan=4,rowspan=4)

root.rowconfigure(0,weight=1)
root.columnconfigure(0,weight=1)


mainframe = tk.Frame(root,bg='#F0F0F0')
text_to_speech = tk.Frame(root,bg='#F0F0F0')
speech_to_text = tk.Frame(root,bg='#F0F0F0')



for frame in (mainframe, text_to_speech, speech_to_text):
    frame.grid(sticky='snew',column=0,row=0,columnspan=3,rowspan=3)

show_frame(mainframe)


#===========mainframe
#logo
logo = Image.open('textP_logo.png')
logo = ImageTk.PhotoImage(logo)
logo_label = tk.Label(mainframe, image=logo,anchor="center")

#place into a column
logo_label.place(relx = 0.5,
                 rely = 0.4,
                 anchor = 'center')

fontStyle = tkFont.Font(family="Lucida Grande", size=18)
instructions = tk.Label(mainframe, text="Free text to speech/speech to text software", font=fontStyle)
instructions.place(relx = 0.5,
                 rely = 0.56,
                 anchor = 'center')

TTS = tk.Button(mainframe,text='text to speech',width=18,height=2,bg="#113245",fg="white",command=lambda:show_frame(text_to_speech))
TTS.place(relx = 0.4,
          rely = 0.7,
          anchor = 'center')
STT = tk.Button(mainframe,text='speech to text',width=18,height=2,bg="#113245",fg="white",command=lambda:show_frame(speech_to_text))
STT.place(relx = 0.6,
          rely = 0.7,
          anchor = 'center')

#===========text_to_speech 

fontStyle = tkFont.Font(family="Lucida Grande", size=18)
instructions_TTS = tk.Label(text_to_speech, text="Text to Speech", font=fontStyle)
instructions_TTS.place(relx = 0.5,
                 rely = 0.05,
                 anchor = 'center')



text_box_TTS = tk.Text(text_to_speech,padx=15,pady=15,bg="white")
#text_box.insert(1.0, "page_content")
#text_box.tag_configure("center", justify="center")
#text_box.tag_add("center",1.0,"end")
#text_box.grid(column=1,row=3)

text_box_TTS.place(relx = 0.5,
                 rely = 0.45,
                 anchor = 'center')


def text_to_speech_func():
    try:
        result=text_box_TTS.get("1.0","end")
        #print(result)
        #print("hello")
        engine = pyttsx3.init() 
        engine.say(result)  
        engine.runAndWait() 
          
    except: 
        print("unknown error occured") 

listen_button_TTS = tk.Button(text_to_speech,text='listen',width=18,height=2,bg="#113245",fg="white",command=lambda:text_to_speech_func())
listen_button_TTS.place(relx = 0.5,
          rely = 0.85,
          anchor = 'center')

def openFile():
    text_box_TTS.insert(1.0, "")
    
    browse_text_TTS.set("loading...")
    file = askopenfile(parent=root, mode='rb', title="Choose a file", filetype=[("word file","*.docx"),("text file","*.txt")])
    
    
    
    if file:
        path = pathlib.Path(file.name)
        text_box_TTS.delete(1.0,"end")
        '''
        if path.suffix == ".pdf":
            print(path.suffix)
            
            read_pdf = PyPDF2.PdfFileReader(file)
        
            page = [0]*read_pdf.getNumPages()
            page_content = ""
            for i in range(read_pdf.getNumPages()):
                page[i] = read_pdf.getPage(i)
                words = page[i].extractText()
                page_content += words
            
            #textbox
            

            text_box.insert(1.0, page_content)

            browse_text.set("browse file")
            '''

        if path.suffix == ".docx":
            
            doc = docx.Document(file.name)
            all_paras = doc.paragraphs
            
            #print("hello")
            doc = docx.Document (file)
            completedText = []
            page_content = ""
            for paragraph in doc.paragraphs:
                completedText.append (paragraph.text)
                #print(paragraph.text) 
                page_content += " "+paragraph.text

            text_box_TTS.insert(1.0, page_content)
            browse_text_TTS.set("browse file")
        
        elif path.suffix == ".txt":
            f = open(file.name, "r")
            text_box_TTS.insert(1.0, f.read())
            browse_text_TTS.set("browse file")
        
        else:
            browse_text_TTS.set("browse file")
    elif file ==None:
        print("hello")
        browse_text_TTS.set("browse file")


browse_Font_TTS = tkFont.Font(size=7)
browse_text_TTS = tk.StringVar()
browse_btn_TTS = tk.Button(text_to_speech,textvariable=browse_text_TTS,bg="#113245",fg="white",command=lambda:openFile(),font=browse_Font_TTS)
browse_text_TTS.set("browse file")
browse_btn_TTS.place(relx = 0.84,
          rely = 0.77,
          anchor = 'center')

back_Font_TTS = tkFont.Font(size=10)
back_btn_TTS = tk.Button(text_to_speech,text="back",bg="white",fg="black",command=lambda:delete_content(),font=back_Font_TTS)
back_btn_TTS.place(relx = 0.048,
          rely = 0.95,
          anchor = 'ne')











#===========speech_to_text
fontStyle = tkFont.Font(family="Lucida Grande", size=18)
instructions_STT = tk.Label(speech_to_text, text="Speech to Text", font=fontStyle)
instructions_STT.place(relx = 0.5,
                 rely = 0.05,
                 anchor = 'center')


text_box_STT = tk.Text(speech_to_text,padx=15,pady=15,bg="white")

text_box_STT.place(relx = 0.5,
                 rely = 0.45,
                 anchor = 'center')


def speech_to_text_func():
    try: 
          
        # use the microphone as source for input. 
        with sr.Microphone() as source2: 
              
            # wait for a second to let the recognizer 
            # adjust the energy threshold based on 
            # the surrounding noise level  
            r.adjust_for_ambient_noise(source2, duration=1) 
              
            #listens for the user's input  
            audio2 = r.listen(source2) 
              
            # Using ggogle to recognize audio 
            MyText = r.recognize_google(audio2) 
            MyText = MyText.lower() 
  
            print("Did you say "+MyText) 
            textToSpeak = text_box_STT.get(1.0,"end")
            text_box_STT.delete(1.0,"end")
            text_box_STT.insert(1.0, textToSpeak +""+MyText)

    except sr.RequestError as e: 
        print("Could not request results; {0}".format(e)) 
          
    except sr.UnknownValueError: 
        print("unknown error occured") 

speak_button_STT = tk.Button(speech_to_text,text='speak',width=18,height=2,bg="#113245",fg="white",command=lambda:speech_to_text_func())
speak_button_STT.place(relx = 0.5,
          rely = 0.85,
          anchor = 'center')

def copy():
    clipboard.copy(text_box_STT.get("1.0","end"))


copy_Font_STT = tkFont.Font(size=7)
copy_text_STT = tk.StringVar()
copy_btn_STT = tk.Button(speech_to_text,textvariable=copy_text_STT,bg="#113245",fg="white",command=lambda:copy(),font=copy_Font_STT)
copy_text_STT.set("copy text")
copy_btn_STT.place(relx = 0.84,
          rely = 0.77,
          anchor = 'center')

back_Font_STT = tkFont.Font(size=10)
back_btn_STT = tk.Button(speech_to_text,text="back",bg="white",fg="black",command=lambda:delete_content(),font=back_Font_STT )
back_btn_STT.place(relx = 0.048,
          rely = 0.95,
          anchor = 'ne')

def delete_content():
    text_box_TTS.delete(1.0,"end")
    text_box_STT.delete(1.0,"end")
    show_frame(mainframe)

root.mainloop()